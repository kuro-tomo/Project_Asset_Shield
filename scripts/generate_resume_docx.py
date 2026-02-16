#!/usr/bin/env python3
"""ProConnect用 職務経歴書 DOCX生成（Google Docs互換）"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT = "/Users/MBP/Desktop/Resume_Quant_Consultant.docx"

NAVY = RGBColor(26, 42, 82)
DARK = RGBColor(33, 33, 33)
GRAY = RGBColor(100, 100, 100)
ACCENT = RGBColor(0, 102, 204)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, title):
    """Navy background section heading."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(17)
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "1A2A52")
    p = cell.paragraphs[0]
    run = p.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.space_after = Pt(0)
    p.space_before = Pt(0)
    doc.add_paragraph()  # spacing


def add_subsection(doc, title):
    """Subsection with navy underline effect."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    p.space_after = Pt(2)
    p.space_before = Pt(6)


def add_kv(doc, key, val):
    """Key-value pair line."""
    p = doc.add_paragraph()
    p.space_after = Pt(1)
    p.space_before = Pt(1)
    run_k = p.add_run(key + "　")
    run_k.bold = True
    run_k.font.size = Pt(9)
    run_k.font.color.rgb = DARK
    run_v = p.add_run(val)
    run_v.font.size = Pt(9)
    run_v.font.color.rgb = DARK


def add_body(doc, text):
    """Body text paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    p.space_after = Pt(3)
    p.space_before = Pt(1)


def add_bullet(doc, text):
    """Bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    p.space_after = Pt(1)
    p.space_before = Pt(1)


def add_skill_row(table, name, level, desc):
    """Add a skill row to a table."""
    row = table.add_row()
    # Name
    c0 = row.cells[0]
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(name)
    r0.font.size = Pt(9)
    r0.font.color.rgb = DARK
    # Bar
    c1 = row.cells[1]
    p1 = c1.paragraphs[0]
    bar = "■" * level + "□" * (5 - level)
    r1 = p1.add_run(bar)
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = NAVY
    # Description
    c2 = row.cells[2]
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(desc)
    r2.font.size = Pt(8)
    r2.font.color.rgb = GRAY


def add_skill_table(doc, skills):
    """Create a skill table section."""
    tbl = doc.add_table(rows=0, cols=3)
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(3)
    tbl.columns[2].width = Cm(9)
    for name, level, desc in skills:
        add_skill_row(tbl, name, level, desc)
    doc.add_paragraph()  # spacing


def add_project(doc, period, title, role, desc_lines, tech):
    """Project block."""
    # Title line
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(2)
    run_period = p.add_run(period + "　　")
    run_period.bold = True
    run_period.font.size = Pt(9)
    run_period.font.color.rgb = NAVY
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = DARK
    # Role
    p_role = doc.add_paragraph()
    p_role.space_after = Pt(2)
    run_role = p_role.add_run(f"役割: {role}")
    run_role.font.size = Pt(8.5)
    run_role.font.color.rgb = GRAY
    # Description bullets
    for line in desc_lines:
        add_bullet(doc, line)
    # Tech stack
    p_tech = doc.add_paragraph()
    p_tech.space_before = Pt(3)
    run_label = p_tech.add_run("技術スタック: ")
    run_label.bold = True
    run_label.font.size = Pt(8)
    run_label.font.color.rgb = ACCENT
    run_tech = p_tech.add_run(tech)
    run_tech.font.size = Pt(8)
    run_tech.font.color.rgb = DARK


def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Hiragino Kaku Gothic Pro"
    style.font.size = Pt(9)

    # ===== TITLE =====
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("職 務 経 歴 書")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_d = p_date.add_run("作成日: 2026年2月10日")
    run_d.font.size = Pt(8)
    run_d.font.color.rgb = GRAY

    # Horizontal line
    p_line = doc.add_paragraph()
    p_line.space_after = Pt(2)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A2A52"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)

    # ===== 基本情報 =====
    add_section_heading(doc, "基本情報")
    add_kv(doc, "氏名", "[要記入]")
    add_kv(doc, "生年月日", "[要記入]")
    add_kv(doc, "所在地", "[要記入]  （例: 東京都）")
    add_kv(doc, "稼働可能時期", "即日（2026年2月〜）")
    add_kv(doc, "稼働希望", "週3〜5日（リモート中心、常駐も可）")
    add_kv(doc, "希望単価", "月額 100〜150万円（税別・スコープにより応相談）")

    # ===== 職務要約 =====
    add_section_heading(doc, "職務要約")
    add_body(
        doc,
        "クオンツ・エンジニア / アルゴリズム・ストラテジスト。"
        "日米両市場を対象としたマルチファクター運用戦略の設計・実装・検証に従事。",
    )
    add_body(
        doc,
        "【米国株】15年バックテスト（2010-2024）で CAGR 14.4%、Sharpe 0.92 を達成した "
        "Core-Satellite型戦略を独力で構築し、現在ライブ運用中。",
    )
    add_body(
        doc,
        "【日本株】J-Quants全銘柄（約5,000銘柄・1,490万レコード）を対象とした "
        "18年間（2008-2026）の大規模バックテスト基盤を構築。"
        "4フェーズのストレステスト（リーマン〜金利正常化）をサバイバーシップバイアスフリーで実施。",
    )
    add_body(
        doc,
        "Python / QuantConnect (LEAN) / Interactive Brokers を用いた"
        "フルスタックのクオンツ開発（データ取得 → ファクター算出 → バックテスト → "
        "ライブ実行 → リスク管理 → 通知）に精通。"
        "金融工学の理論とソフトウェアエンジニアリングの実務を両立できる稀少な人材。",
    )

    # ===== 技術スキル =====
    add_section_heading(doc, "技術スキル")

    add_subsection(doc, "プログラミング言語")
    add_skill_table(doc, [
        ("Python", 5, "10年+（NumPy/Pandas/SciPy/scikit-learn）"),
        ("SQL", 4, "8年+（SQLite/PostgreSQL）"),
        ("C# / .NET", 3, "3年+（QuantConnect LEAN）"),
        ("JavaScript/TypeScript", 3, "3年+"),
    ])

    add_subsection(doc, "クオンツ・金融専門")
    add_skill_table(doc, [
        ("ファクターモデル設計", 5, "マルチファクター（5因子）、ランク正規化"),
        ("バックテスト設計", 5, "Walk-Forward、IS/OOS分割、過学習防止"),
        ("リスク管理", 5, "Kill-switch、逆ボラサイジング、DD制御"),
        ("レジーム判定", 5, "マクロ指標ベースの相場局面分類"),
        ("ポートフォリオ構築", 5, "Core-Satellite、セクター中立化"),
        ("統計分析", 4, "仮説検定、ブートストラップ、モンテカルロ"),
    ])

    add_subsection(doc, "プラットフォーム / インフラ")
    add_skill_table(doc, [
        ("QuantConnect (LEAN)", 5, "アルゴ設計〜ライブ運用まで一貫対応"),
        ("Interactive Brokers API", 4, "注文管理、ポジション管理、データ取得"),
        ("Git / GitHub", 5, "バージョン管理、CI/CD"),
        ("Linux / macOS", 4, "サーバー構築、スクリプト自動化"),
        ("Docker", 3, "コンテナ化、再現可能環境"),
        ("AWS / GCP", 3, "EC2/S3、Cloud Functions"),
    ])

    add_subsection(doc, "データ基盤")
    add_skill_table(doc, [
        ("金融データパイプライン", 5, "ETL設計、データ品質検証"),
        ("SQLite / PostgreSQL", 4, "大規模時系列データ管理"),
        ("J-Quants API", 4, "日本株データ取得・加工"),
        ("REST API設計", 4, "認証、ページネーション、レート制御"),
    ])

    # ===== プロジェクト実績 =====
    add_section_heading(doc, "プロジェクト実績")

    add_project(
        doc,
        "2023 〜 現在",
        "米国株 Core-Satellite 運用戦略（V8）",
        "設計・開発・運用（個人プロジェクト）",
        [
            "SPY 80% Core + 50% マルチファクター Satellite の 130% Core-Satellite 構造を設計",
            "5ファクターモデル（Momentum / Short-term Momentum / Low Volatility / Value / Quality）を構築",
            "マクロ指標ベースのレジーム判定（5段階）で動的にエクスポージャーを調整",
            "DD 15%で自動発動する Kill-switch を実装（Satellite清算、Core維持）",
            "15年バックテスト（2010-2024）: リターン +649%、CAGR 14.4%、Sharpe 0.92、Max DD 25.7%",
            "Walk-Forward検証: IS Sharpe 1.10 / OOS Sharpe 0.74（頑健性を確認）",
            "Email通知システム（Kill-switch発動、レジーム変更、週次サマリー）を実装",
            "QuantConnect上でペーパートレード → Interactive Brokers ライブ移行を推進中",
        ],
        "Python, QuantConnect (LEAN/C#), Interactive Brokers API, NumPy, Pandas, SciPy, SQLite",
    )

    add_project(
        doc,
        "2022 〜 2023",
        "マルチファクター戦略 研究開発（V1〜V7）",
        "設計・開発・検証（個人プロジェクト）",
        [
            "米国株式のファクター投資戦略を7世代にわたり反復改善",
            "V6 Pro（純粋ファクター Long-Only）がSPYを下回る課題を特定 → Core-Satellite構造（V8）で解決",
            "パラメータチューニングだけでは構造的問題を解決できないことを実証",
            "逆ボラティリティ・サイジング、セクター分散制約、トレーリングストップ等のリスク管理手法を体系化",
            "41モジュール構成のPythonライブラリ（src/shield/）として知見をコード化",
        ],
        "Python, QuantConnect, pandas, scikit-learn, matplotlib, SQLite",
    )

    add_project(
        doc,
        "2021 〜 現在",
        "日本株 全銘柄バックテスト基盤（18年・5,000銘柄）",
        "設計・開発・検証（個人プロジェクト）",
        [
            "J-Quants API の全銘柄データ（約5,000銘柄、1,490万レコード）を用いた大規模バックテスト基盤を構築",
            "サバイバーシップバイアスフリー: 上場廃止銘柄を含む全銘柄で検証可能な設計",
            "4フェーズ・ストレステスト: リーマンショック → アベノミクス → COVID → 金利正常化",
            "Walk-Forward検証フレームワーク（IS/Validation/OOS 3分割）を実装",
            "SQLiteデータウェアハウス（2.3GB）、データ品質カバレッジ99%+を確認",
            "41モジュール構成のPython量的分析ライブラリを構築",
        ],
        "Python, J-Quants API, SQLite, NumPy, Pandas, SciPy, scikit-learn, REST API",
    )

    add_project(
        doc,
        "2024 〜 現在",
        "金融データパイプライン基盤構築",
        "設計・開発（個人プロジェクト）",
        [
            "J-Quants API を用いた日本株全銘柄のデータ取得・加工パイプラインを構築",
            "固定長バイナリデータ（374バイト、cp932）のパーサーを開発、83フィールドの自動抽出を実現",
            "SQLiteベースのデータウェアハウス設計（1,490万レコード規模）",
            "データ品質検証フレームワーク（カバレッジ率99%+を確認）を実装",
            "ETLパイプライン: ダウンローダー → パーサー → ETL → スキーマ → SQLite の自動化",
        ],
        "Python, SQLite, J-Quants API, REST API, cp932エンコーディング, ETL設計",
    )

    # ===== 強み・自己PR =====
    add_section_heading(doc, "強み・自己PR")

    add_subsection(doc, "1. 戦略設計から実運用まで一気通貫")
    add_body(
        doc,
        "アイデア着想 → 数理モデル設計 → Pythonコード実装 → バックテスト検証 → "
        "リスク管理設計 → ライブ運用基盤構築 → モニタリング・通知の全工程を一人で完遂できます。"
        "外注や分業なしで、構想から本番稼働まで最短距離で到達可能です。",
    )

    add_subsection(doc, "2. 失敗から学ぶ反復改善力")
    add_body(
        doc,
        "7世代の戦略改善を通じて「パラメータチューニングでは構造的問題は解決できない」"
        "「ファクター Long-Only はベータリスクに脆弱」等の本質的な教訓を獲得。"
        "V6の失敗（SPY下回り）を分析し、Core-Satellite構造（V8）で根本解決に至りました。"
        "この反復改善プロセスは、クライアントの既存戦略の改善にも直接応用できます。",
    )

    add_subsection(doc, "3. データ品質への執着")
    add_body(
        doc,
        "固定長バイナリ（cp932/374バイト）の完全パーサー開発、83フィールドの自動検証、"
        "99%+のカバレッジ率確認など、データ品質に対して妥協しません。"
        "「ガベージイン・ガベージアウト」を防ぐデータパイプラインの設計・運用が可能です。",
    )

    add_subsection(doc, "4. 再現可能な研究プロセス")
    add_body(
        doc,
        "Walk-Forward検証（IS/OOS分割）、決定論的バックテスト、Git管理による完全な再現性を確保。"
        "過学習・サバイバーシップバイアス・ルックアヘッドバイアス等の定量リサーチの落とし穴を熟知し、"
        "堅牢な検証フレームワークを構築できます。",
    )

    # ===== 対応可能な案件例 =====
    add_section_heading(doc, "対応可能な案件例")
    add_bullet(doc, "クオンツ戦略の設計・開発・バックテスト支援")
    add_bullet(doc, "ファクターモデル / リスクモデルの構築")
    add_bullet(doc, "バックテストフレームワークの設計・開発")
    add_bullet(doc, "金融データパイプライン（ETL）の構築")
    add_bullet(doc, "QuantConnect / Interactive Brokers を用いた自動売買システム開発")
    add_bullet(doc, "既存クオンツ戦略のレビュー・改善提案")
    add_bullet(doc, "Pythonによる金融データ分析基盤の構築")

    # Footer note
    doc.add_paragraph()
    p_note1 = doc.add_paragraph()
    p_note1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_note1.add_run(
        "※ 本経歴書に記載の戦略パフォーマンス指標はバックテスト結果であり、将来の成果を保証するものではありません。"
    )
    r1.font.size = Pt(7.5)
    r1.font.color.rgb = GRAY

    p_note2 = doc.add_paragraph()
    p_note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_note2.add_run(
        "※ 戦略の具体的ロジック・パラメータは営業秘密として非開示とさせていただきます。"
    )
    r2.font.size = Pt(7.5)
    r2.font.color.rgb = GRAY

    doc.save(OUTPUT)
    print(f"DOCX generated: {OUTPUT}")


if __name__ == "__main__":
    build()
