#!/usr/bin/env python3
"""English Resume DOCX for international quant roles (Google Docs compatible)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

OUTPUT = "/Users/MBP/Desktop/Resume_Quant_Engineer_EN.docx"

NAVY = RGBColor(26, 42, 82)
DARK = RGBColor(33, 33, 33)
GRAY = RGBColor(100, 100, 100)
ACCENT = RGBColor(0, 102, 204)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, title):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(17)
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "1A2A52")
    p = cell.paragraphs[0]
    run = p.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.space_after = Pt(0)
    p.space_before = Pt(0)
    doc.add_paragraph()


def add_subsection(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    p.space_after = Pt(2)
    p.space_before = Pt(6)


def add_kv(doc, key, val):
    p = doc.add_paragraph()
    p.space_after = Pt(1)
    p.space_before = Pt(1)
    run_k = p.add_run(key + "  ")
    run_k.bold = True
    run_k.font.size = Pt(9)
    run_k.font.color.rgb = DARK
    run_v = p.add_run(val)
    run_v.font.size = Pt(9)
    run_v.font.color.rgb = DARK


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    p.space_after = Pt(3)
    p.space_before = Pt(1)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    p.space_after = Pt(1)
    p.space_before = Pt(1)


def add_skill_row(table, name, level, desc):
    row = table.add_row()
    c0 = row.cells[0]
    r0 = c0.paragraphs[0].add_run(name)
    r0.font.size = Pt(9)
    r0.font.color.rgb = DARK
    c1 = row.cells[1]
    bar = "■" * level + "□" * (5 - level)
    r1 = c1.paragraphs[0].add_run(bar)
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = NAVY
    c2 = row.cells[2]
    r2 = c2.paragraphs[0].add_run(desc)
    r2.font.size = Pt(8)
    r2.font.color.rgb = GRAY


def add_skill_table(doc, skills):
    tbl = doc.add_table(rows=0, cols=3)
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(3)
    tbl.columns[2].width = Cm(9)
    for name, level, desc in skills:
        add_skill_row(tbl, name, level, desc)
    doc.add_paragraph()


def add_project(doc, period, title, role, desc_lines, tech):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(2)
    run_period = p.add_run(period + "    ")
    run_period.bold = True
    run_period.font.size = Pt(9)
    run_period.font.color.rgb = NAVY
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = DARK
    p_role = doc.add_paragraph()
    p_role.space_after = Pt(2)
    run_role = p_role.add_run(f"Role: {role}")
    run_role.font.size = Pt(8.5)
    run_role.font.color.rgb = GRAY
    for line in desc_lines:
        add_bullet(doc, line)
    p_tech = doc.add_paragraph()
    p_tech.space_before = Pt(3)
    run_label = p_tech.add_run("Tech Stack: ")
    run_label.bold = True
    run_label.font.size = Pt(8)
    run_label.font.color.rgb = ACCENT
    run_tech = p_tech.add_run(tech)
    run_tech.font.size = Pt(8)
    run_tech.font.color.rgb = DARK


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # ===== HEADER =====
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_name.add_run("[YOUR NAME]")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("Quantitative Engineer / Algorithmic Strategist")
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = GRAY

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.space_after = Pt(2)
    run_c = p_contact.add_run("[Location]  |  [Email]  |  [LinkedIn URL]")
    run_c.font.size = Pt(9)
    run_c.font.color.rgb = GRAY

    # Horizontal line
    p_line = doc.add_paragraph()
    p_line.space_after = Pt(2)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A2A52"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)

    # ===== SUMMARY =====
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    add_body(
        doc,
        "Quantitative engineer and algorithmic strategist specializing in systematic "
        "equity strategies across U.S. and Japanese markets.",
    )
    add_body(
        doc,
        "[U.S. Equity] Independently designed and built a Core-Satellite multi-factor strategy "
        "achieving 14.4% CAGR and Sharpe Ratio 0.92 over a 15-year backtest period (2010-2024), "
        "validated through rigorous Walk-Forward analysis (IS Sharpe 1.10 / OOS Sharpe 0.74). "
        "Currently live on QuantConnect + Interactive Brokers.",
    )
    add_body(
        doc,
        "[Japanese Equity] Built large-scale backtesting infrastructure covering ~5,000 stocks "
        "(14.9M records) over 18 years (2008-2026) via J-Quants API. Survivorship-bias-free "
        "4-phase stress test: Lehman Crisis, Abenomics, COVID, Rate Normalization.",
    )
    add_body(
        doc,
        "Full-stack quant development expertise spanning the entire pipeline: "
        "data acquisition, factor construction, backtesting, risk management, "
        "live execution (QuantConnect/LEAN + Interactive Brokers), and monitoring. "
        "Rare combination of quantitative finance theory and production-grade "
        "software engineering.",
    )

    # ===== KEY METRICS =====
    add_section_heading(doc, "STRATEGY PERFORMANCE HIGHLIGHTS")
    tbl = doc.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics = [
        ("Cumulative Return", "+649%"),
        ("CAGR", "14.4%"),
        ("Sharpe Ratio", "0.92"),
        ("Max Drawdown", "25.7%"),
        ("IS Sharpe", "1.10"),
        ("OOS Sharpe", "0.74"),
        ("Backtest Period", "15 years"),
        ("Walk-Forward", "Validated"),
    ]
    for i, (label, value) in enumerate(metrics):
        row_idx = i // 4
        col_idx = i % 4
        cell = tbl.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_val = p.add_run(value + "\n")
        r_val.bold = True
        r_val.font.size = Pt(11)
        r_val.font.color.rgb = NAVY
        r_lab = p.add_run(label)
        r_lab.font.size = Pt(7.5)
        r_lab.font.color.rgb = GRAY
    doc.add_paragraph()

    # ===== TECHNICAL SKILLS =====
    add_section_heading(doc, "TECHNICAL SKILLS")

    add_subsection(doc, "Programming Languages")
    add_skill_table(doc, [
        ("Python", 5, "10+ yrs (NumPy, Pandas, SciPy, scikit-learn)"),
        ("SQL", 4, "8+ yrs (SQLite, PostgreSQL)"),
        ("C# / .NET", 3, "3+ yrs (QuantConnect LEAN engine)"),
        ("JavaScript / TypeScript", 3, "3+ yrs"),
    ])

    add_subsection(doc, "Quantitative Finance")
    add_skill_table(doc, [
        ("Factor Model Design", 5, "Multi-factor (5-factor), rank normalization"),
        ("Backtesting", 5, "Walk-Forward, IS/OOS split, overfitting prevention"),
        ("Risk Management", 5, "Kill-switch, inverse-vol sizing, DD control"),
        ("Regime Detection", 5, "Macro-indicator-based market regime classification"),
        ("Portfolio Construction", 5, "Core-Satellite, sector neutralization"),
        ("Statistical Analysis", 4, "Hypothesis testing, bootstrap, Monte Carlo"),
    ])

    add_subsection(doc, "Platforms & Infrastructure")
    add_skill_table(doc, [
        ("QuantConnect (LEAN)", 5, "End-to-end: algo design to live trading"),
        ("Interactive Brokers API", 4, "Order mgmt, position mgmt, market data"),
        ("Git / GitHub", 5, "Version control, CI/CD"),
        ("Linux / macOS", 4, "Server setup, scripting, automation"),
        ("Docker", 3, "Containerization, reproducible environments"),
        ("AWS / GCP", 3, "EC2/S3, Cloud Functions"),
    ])

    add_subsection(doc, "Data Engineering")
    add_skill_table(doc, [
        ("Financial Data Pipelines", 5, "ETL design, data quality validation"),
        ("SQLite / PostgreSQL", 4, "Large-scale time-series data management"),
        ("REST API Integration", 4, "Auth, pagination, rate limiting"),
        ("J-Quants API", 4, "Japanese equity data acquisition"),
    ])

    # ===== PROJECTS =====
    add_section_heading(doc, "PROJECT EXPERIENCE")

    add_project(
        doc,
        "2023 - Present",
        "U.S. Equity Core-Satellite Strategy (V8)",
        "Sole Designer, Developer & Operator",
        [
            "Designed 130% Core-Satellite structure: 80% SPY core + 50% multi-factor satellite",
            "Built 5-factor model: Momentum, Short-term Momentum, Low Volatility, Value (P/E), Quality (ROE)",
            "Implemented 5-stage macro-based regime detection for dynamic exposure adjustment",
            "Developed automatic kill-switch: triggers at 15% DD, liquidates satellite, preserves core",
            "15-year backtest (2010-2024): +649% return, 14.4% CAGR, 0.92 Sharpe, 25.7% max DD",
            "Walk-Forward validation: IS Sharpe 1.10 / OOS Sharpe 0.74 (robustness confirmed)",
            "Built email alert system for kill-switch triggers, regime changes, and weekly summaries",
            "Currently transitioning from paper trading to Interactive Brokers live execution",
        ],
        "Python, QuantConnect (LEAN/C#), Interactive Brokers API, NumPy, Pandas, SciPy, SQLite",
    )

    add_project(
        doc,
        "2022 - 2023",
        "Multi-Factor Strategy R&D (V1-V7)",
        "Sole Designer & Developer",
        [
            "Iterated through 7 generations of U.S. equity factor strategies",
            "Identified structural flaw in V6 (pure factor long-only underperformed SPY) and resolved with Core-Satellite architecture (V8)",
            "Proved that parameter tuning alone cannot fix structural portfolio problems",
            "Systematized risk management: inverse-vol sizing, sector diversification constraints, trailing stops",
            "Codified research into a 41-module Python library (src/shield/)",
        ],
        "Python, QuantConnect, pandas, scikit-learn, matplotlib, SQLite",
    )

    add_project(
        doc,
        "2021 - Present",
        "Japanese Equity Backtesting Infrastructure (18yr / 5,000 Stocks)",
        "Sole Designer & Developer",
        [
            "Built large-scale backtesting infrastructure using J-Quants API (~5,000 stocks, 14.9M records)",
            "Survivorship-bias-free design: includes delisted stocks for unbiased testing",
            "4-phase stress test: Lehman Crisis (2006-2010), Abenomics (2011-2015), COVID (2016-2020), Rate Normalization (2021-2026)",
            "Implemented Walk-Forward validation framework (IS / Validation / OOS 3-way split)",
            "SQLite data warehouse (2.3GB), data quality coverage 99%+",
            "Built 41-module Python quantitative analysis library",
        ],
        "Python, J-Quants API, SQLite, NumPy, Pandas, SciPy, scikit-learn, REST API",
    )

    add_project(
        doc,
        "2024 - Present",
        "Financial Data Pipeline Engineering",
        "Sole Designer & Developer",
        [
            "Built end-to-end data pipeline for all Japanese equities via J-Quants API",
            "Developed parser for fixed-width binary data (374-byte, cp932 encoding, 83 fields)",
            "Designed SQLite data warehouse (14.9M records)",
            "Implemented data quality framework achieving 99%+ coverage across all fields",
            "Automated full ETL: downloader -> parser -> transform -> schema -> SQLite",
        ],
        "Python, SQLite, J-Quants API, REST API, ETL Design",
    )

    # ===== STRENGTHS =====
    add_section_heading(doc, "KEY STRENGTHS")

    add_subsection(doc, "1. End-to-End Ownership")
    add_body(
        doc,
        "Capable of single-handedly executing the full quant workflow: "
        "ideation, mathematical modeling, Python implementation, backtesting, "
        "risk management design, live infrastructure, and monitoring. "
        "Zero dependency on external teams to go from concept to production.",
    )

    add_subsection(doc, "2. Iterative Improvement Through Failure")
    add_body(
        doc,
        "Through 7 generations of strategy refinement, extracted fundamental lessons: "
        "'parameter tuning cannot fix structural problems' and 'factor long-only is vulnerable "
        "to beta risk.' Analyzed V6 underperformance vs. SPY and architected the Core-Satellite "
        "solution (V8). This iterative process directly applies to improving client strategies.",
    )

    add_subsection(doc, "3. Data Quality Obsession")
    add_body(
        doc,
        "Built a complete parser for fixed-width binary data (cp932, 374-byte records), "
        "automated validation of 83 fields, and confirmed 99%+ coverage. "
        "Uncompromising on data integrity to prevent garbage-in-garbage-out failures.",
    )

    add_subsection(doc, "4. Reproducible Research")
    add_body(
        doc,
        "Walk-Forward validation (IS/OOS split), deterministic backtesting, and full Git version control "
        "ensure complete reproducibility. Deep awareness of quantitative research pitfalls: "
        "overfitting, survivorship bias, and look-ahead bias.",
    )

    # ===== ENGAGEMENT TYPES =====
    add_section_heading(doc, "AVAILABLE FOR")
    add_bullet(doc, "Quantitative strategy design, development, and backtesting")
    add_bullet(doc, "Factor model / risk model construction")
    add_bullet(doc, "Backtesting framework architecture and development")
    add_bullet(doc, "Financial data pipeline (ETL) engineering")
    add_bullet(doc, "Algorithmic trading systems (QuantConnect / Interactive Brokers)")
    add_bullet(doc, "Quant strategy review and optimization consulting")
    add_bullet(doc, "Python-based financial analytics infrastructure")

    # Footer
    doc.add_paragraph()
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_note.add_run(
        "Note: Strategy performance figures are based on backtested results and do not guarantee future performance. "
        "Specific strategy logic and parameters are proprietary and confidential."
    )
    r1.font.size = Pt(7.5)
    r1.font.color.rgb = GRAY

    doc.save(OUTPUT)
    print(f"DOCX generated: {OUTPUT}")


if __name__ == "__main__":
    build()
